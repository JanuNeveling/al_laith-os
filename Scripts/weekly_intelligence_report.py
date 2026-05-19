#!/usr/bin/env python3
"""
Al Laith Weekly Marketing Intelligence Report
Runs every Monday — researches competitor social media in real time,
then generates a formatted .docx on the Desktop ready to upload to Google Docs.

APIs used:
  - Tavily  (free tier — web + social search)  → tavily.com
  - Anthropic Claude (analysis + writing)       → console.anthropic.com
"""

import os
import sys
from datetime import datetime

# ── Config ──────────────────────────────────────────────────────────────────
CLAUDE_MODEL = 'claude-haiku-4-5'   # ~$0.08/run. Swap to claude-sonnet-4-5 for richer prose (~$0.35/run)
OUTPUT_DIR   = os.path.join(os.path.expanduser('~'), 'Desktop')

COMPANY = {
    'name'      : 'Al Laith Group',
    'sector'    : 'equipment rental, scaffolding, and events infrastructure',
    'market'    : 'UAE (Dubai, Abu Dhabi, wider GCC)',
    'website'   : 'https://www.allaith.com',
    'linkedin'  : 'https://www.linkedin.com/company/al-laith',
    'instagram' : 'https://www.instagram.com/allaithgroup',
}

# Add social handles as you find them — leave blank if unknown
COMPETITORS = [
    {
        'name'      : 'Byrne Equipment Rental',
        'sector'    : 'Equipment Rental',
        'website'   : 'https://www.byrne.ae',
        'linkedin'  : 'https://www.linkedin.com/company/byrne-equipment-rental',
        'instagram' : 'https://www.instagram.com/byrne_equipment',
        'facebook'  : 'https://www.facebook.com/ByrneEquipmentRental',
    },
    {
        'name'      : 'All Events Services (AES)',
        'sector'    : 'Events Infrastructure',
        'website'   : 'https://www.alleventsservices.com',
        'linkedin'  : 'https://www.linkedin.com/company/all-events-services',
        'instagram' : 'https://www.instagram.com/alleventsservices',
        'facebook'  : '',
    },
    {
        'name'      : 'Altrad Group',
        'sector'    : 'Scaffolding & Industrial Services',
        'website'   : 'https://www.altrad.com',
        'linkedin'  : 'https://www.linkedin.com/company/altrad',
        'instagram' : 'https://www.instagram.com/altradgroup',
        'facebook'  : '',
    },
    {
        'name'      : 'Brandsafway',
        'sector'    : 'Scaffolding & Access',
        'website'   : 'https://www.brandsafway.com',
        'linkedin'  : 'https://www.linkedin.com/company/brandsafway',
        'instagram' : 'https://www.instagram.com/brandsafway',
        'facebook'  : '',
    },
    {
        'name'      : 'Alimak',
        'sector'    : 'Vertical Access & Hoists',
        'website'   : 'https://www.alimak.com',
        'linkedin'  : 'https://www.linkedin.com/company/alimak-group',
        'instagram' : 'https://www.instagram.com/alimakgroup',
        'facebook'  : '',
    },
    {
        'name'      : 'Manlift',
        'sector'    : 'Aerial Work Platforms & Access Equipment',
        'website'   : 'https://www.manlift.ae',
        'linkedin'  : 'https://www.linkedin.com/company/manlift-middle-east',
        'instagram' : 'https://www.instagram.com/manlift_me',
        'facebook'  : '',
    },
    {
        'name'      : 'Access Rental Gulf',
        'sector'    : 'Access Equipment Rental',
        'website'   : 'https://www.accessrentalgulf.com',
        'linkedin'  : 'https://www.linkedin.com/company/access-rental-gulf',
        'instagram' : '',
        'facebook'  : '',
    },
    {
        'name'      : 'Johnson Arabia',
        'sector'    : 'Industrial Services & Equipment',
        'website'   : 'https://www.johnsonarabia.com',
        'linkedin'  : 'https://www.linkedin.com/company/johnson-arabia',
        'instagram' : '',
        'facebook'  : '',
    },
    {
        'name'      : 'Arena',
        'sector'    : 'Events Infrastructure',
        'website'   : 'https://www.arena.ae',
        'linkedin'  : 'https://www.linkedin.com/company/arena-middle-east',
        'instagram' : 'https://www.instagram.com/arenagroupme',
        'facebook'  : '',
    },
]


# ── API key loaders ──────────────────────────────────────────────────────────
def _load_key(env_name, filename):
    """Read API key from environment variable, then from a local file."""
    key = os.environ.get(env_name, '').strip()
    if key:
        return key
    key_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    if os.path.exists(key_file):
        with open(key_file, 'r') as f:
            key = f.read().strip()
        if key:
            return key
    return None

def get_anthropic_key():
    key = _load_key('ANTHROPIC_API_KEY', '.api_key')
    if not key:
        print('\nERROR: Anthropic API key not found.')
        print('Fix: Create Scripts/.api_key with your sk-ant-... key, or set ANTHROPIC_API_KEY env var.')
        print('See SETUP.md for instructions.\n')
        sys.exit(1)
    return key

def get_tavily_key():
    key = _load_key('TAVILY_API_KEY', '.tavily_key')
    if not key:
        print('\nWARNING: Tavily API key not found — skipping live social media research.')
        print('Fix: Create Scripts/.tavily_key with your Tavily key (free at tavily.com).')
        print('The report will still run using Claude\'s general knowledge instead.\n')
    return key


# ── Tavily search ────────────────────────────────────────────────────────────
def tavily_search(tavily_client, query, max_results=5):
    """Run a Tavily search, return a clean text block of results."""
    if not tavily_client:
        return ''
    try:
        results = tavily_client.search(
            query=query,
            search_depth='advanced',
            max_results=max_results,
            include_answer=True,
        )
        lines = []
        if results.get('answer'):
            lines.append(f'Summary: {results["answer"]}')
            lines.append('')
        for r in results.get('results', []):
            lines.append(f'Source: {r.get("url", "")}')
            lines.append(f'Title: {r.get("title", "")}')
            lines.append(r.get('content', '')[:500])
            lines.append('')
        return '\n'.join(lines).strip()
    except Exception as e:
        print(f'         Tavily search failed for "{query[:50]}...": {e}')
        return ''

def research_competitor_social(tavily_client, comp):
    """
    Run 3 searches per competitor:
      1. General recent marketing / news
      2. LinkedIn activity
      3. Instagram / social media content
    Returns a combined text block.
    """
    if not tavily_client:
        return ''

    year = datetime.now().year
    name = comp['name']
    sector = comp['sector']

    queries = [
        f'"{name}" UAE marketing campaigns news {year}',
        f'"{name}" LinkedIn posts content {year} UAE {sector}',
        f'"{name}" Instagram social media posts content {year}',
    ]

    # If we have specific profile URLs, add targeted searches
    if comp.get('linkedin'):
        queries.append(f'site:linkedin.com "{name}" {year}')
    if comp.get('instagram'):
        queries.append(f'site:instagram.com "{name}" {year}')

    blocks = []
    for q in queries[:4]:  # cap at 4 searches per competitor
        result = tavily_search(tavily_client, q, max_results=3)
        if result:
            blocks.append(f'--- Search: {q} ---\n{result}')

    return '\n\n'.join(blocks)

def research_trends_social(tavily_client):
    """Search for current UAE market trends and global B2B marketing news."""
    if not tavily_client:
        return '', ''

    year  = datetime.now().year
    month = datetime.now().strftime('%B')

    uae_queries = [
        f'UAE equipment rental scaffolding events marketing trends {month} {year}',
        f'UAE B2B construction LinkedIn marketing {year}',
        f'UAE industrial companies social media marketing {year}',
    ]

    b2b_queries = [
        f'B2B marketing trends LinkedIn Instagram {month} {year}',
        f'B2B social media marketing strategy {year} what is working',
        f'B2B content marketing case studies {year}',
    ]

    uae_blocks = []
    for q in uae_queries:
        r = tavily_search(tavily_client, q, max_results=3)
        if r:
            uae_blocks.append(f'--- {q} ---\n{r}')

    b2b_blocks = []
    for q in b2b_queries:
        r = tavily_search(tavily_client, q, max_results=3)
        if r:
            b2b_blocks.append(f'--- {q} ---\n{r}')

    return '\n\n'.join(uae_blocks), '\n\n'.join(b2b_blocks)


# ── Claude call ──────────────────────────────────────────────────────────────
def ask(client, prompt, max_tokens=2000):
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=max_tokens,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return msg.content[0].text


# ── Prompts ──────────────────────────────────────────────────────────────────
def prompt_competitor(comp, social_data=''):
    website_note = f'Website: {comp["website"]}' if comp.get('website') else 'No public website listed.'
    linkedin_note = f'LinkedIn: {comp["linkedin"]}' if comp.get('linkedin') else ''
    instagram_note = f'Instagram: {comp["instagram"]}' if comp.get('instagram') else ''

    social_context = ''
    if social_data:
        social_context = f"""
LIVE SOCIAL MEDIA & WEB RESEARCH (gathered this week):
{social_data[:3000]}

Use the above real data as the primary source for your analysis. Quote or reference specific
posts, campaigns, or content themes where the data supports it.
"""
    else:
        social_context = '\nNo live social data available — base analysis on general knowledge.\n'

    return f"""You are a senior B2B marketing analyst specialising in the UAE construction,
infrastructure, and events sector.

Analyse the marketing strategy of **{comp['name']}** ({comp['sector']}, UAE market).
{website_note}
{linkedin_note}
{instagram_note}
{social_context}
Context: Al Laith Group is a UAE-based equipment rental, scaffolding, and events infrastructure
company competing in the same space.

Write a structured marketing intelligence report using EXACTLY these section headers:

## {comp['name']} — Marketing Intelligence

### Brand Positioning
How they position themselves in the market. Key messages and value propositions.

### Social Media Presence
Which platforms they are active on. Posting frequency, content types, tone of voice.
Reference specific post themes or campaigns from the research data where available.

### Content Strategy
What content themes they lead with (project showcase, safety, team, promotions, etc).
What is working well. What format dominates (video, image, carousel, text).

### Audience & Messaging
Who they target. How they speak to that audience. Key phrases or hooks they use.

### Marketing Strengths
2-3 specific things they do well. Be concrete — reference actual content where possible.

### Marketing Weaknesses & Gaps
Where their marketing falls short. Where Al Laith could clearly differentiate.

### Threat Level for Al Laith
Low / Medium / High — one sentence explanation.

Rules: 3-5 sentences per section. Plain business English. Direct and specific.
If social data is available, prioritise it over general knowledge."""


def prompt_uae_trends(live_data=''):
    month_year = datetime.now().strftime('%B %Y')
    live_context = ''
    if live_data:
        live_context = f"""
LIVE WEB RESEARCH — UAE MARKET (gathered this week):
{live_data[:2500]}

Use this real data as your primary source. Reference specific trends, companies, or posts
from the research where relevant.
"""

    return f"""You are a senior B2B marketing analyst with deep expertise in the UAE and GCC market.

Write a comprehensive marketing trends report for the equipment rental, scaffolding, and events
infrastructure sector in the UAE for {month_year}.
{live_context}
Use EXACTLY these section headers:

## UAE Market Marketing Trends — {month_year}

### Sector Overview
Current state of B2B marketing in UAE equipment rental, scaffolding, and events infrastructure.

### Social Media Landscape
Which platforms are UAE B2B industrial companies most active on right now.
What types of content are getting traction. Reference real examples where the data allows.

### LinkedIn in the UAE
How UAE construction and industrial companies are using LinkedIn.
Post formats, messaging styles, engagement patterns.

### Instagram & Visual Platforms
How this sector uses Instagram, YouTube Shorts, and TikTok.
What visual content approaches are working.

### Buyer Behaviour & Digital Shifts
How B2B procurement is changing in UAE construction and events in 2025.

### Language & Cultural Considerations
Arabic vs English content. Cultural sensitivities. Relationship-driven marketing nuances.

### Top 4 Opportunities for Al Laith Right Now
Four specific, actionable marketing opportunities Al Laith should move on this week.

Rules: Specific and actionable. For a marketing manager making real decisions this week."""


def prompt_global_b2b(live_data=''):
    month_year = datetime.now().strftime('%B %Y')
    live_context = ''
    if live_data:
        live_context = f"""
LIVE WEB RESEARCH — GLOBAL B2B MARKETING (gathered this week):
{live_data[:2500]}

Use this real data as your primary source. Reference specific trends, companies, or campaigns
from the research where relevant.
"""

    return f"""You are a global B2B marketing strategist writing a weekly briefing for the marketing
team of Al Laith Group — a UAE equipment rental, scaffolding, and events infrastructure company.

Write a global B2B marketing intelligence briefing for {month_year}.
{live_context}
Use EXACTLY these section headers:

## Global B2B Marketing Trends — {month_year}

### The Biggest Shift Right Now
The single most significant change happening in global B2B marketing this quarter.

### LinkedIn & Social Selling
What leading B2B companies are doing on LinkedIn that is working right now.
Specific tactics, content formats, posting cadence, engagement strategies.

### Instagram & Short-Form Video for B2B
How B2B companies globally are using Instagram Reels, TikTok, and YouTube Shorts.
What works, what doesn't, which industries are leading.

### Content Marketing That's Actually Converting
Formats, lengths, and approaches winning in B2B content globally.

### AI Tools in B2B Marketing
How B2B teams are using AI practically — content creation, research, personalisation.

### Case Studies & Social Proof
How top B2B companies use project showcases and client testimonials to win business.

### What Al Laith Should Steal This Week
5 specific, implementable tactics from global B2B leaders that Al Laith can use in UAE now.
Be extremely concrete — name the tactic, the platform, and exactly how to execute it.

Rules: Name real companies or campaigns where possible. Concrete, not theoretical."""


def prompt_exec_summary(comp_summaries, uae_snippet, b2b_snippet):
    return f"""You are a Chief Marketing Officer writing a weekly intelligence briefing for Al Laith Group.

Based on this week's research below, write a sharp executive summary — exactly 6 bullet points.

COMPETITOR RESEARCH:
{comp_summaries}

UAE MARKET SNAPSHOT:
{uae_snippet}

GLOBAL B2B SNAPSHOT:
{b2b_snippet}

Rules:
- Exactly 6 bullet points
- Each bullet = one punchy, specific insight or action item (max 28 words)
- Start each bullet with a relevant emoji
- Mix: 2 competitor insights, 2 trend insights, 2 concrete action recommendations for Al Laith
- No headers, no intro text — just the 6 bullets
- Reference specific companies, platforms, or tactics from the research where possible"""


# ── Word document builder ────────────────────────────────────────────────────
def build_document(week_label, exec_summary, competitor_reports, uae_trends, global_b2b, used_live_data):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print('\nERROR: python-docx not installed. Run: pip install python-docx\n')
        sys.exit(1)

    NAVY  = RGBColor(0x1A, 0x2B, 0x4A)
    GOLD  = RGBColor(0xC9, 0x9E, 0x3A)
    GREY  = RGBColor(0x66, 0x66, 0x66)
    LGREY = RGBColor(0x99, 0x99, 0x99)
    GREEN = RGBColor(0x16, 0xA3, 0x4A)

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Cover page ──────────────────────────────────────────────
    def cpara(text, size, bold=False, color=NAVY, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(text)
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.color.rgb = color
        r.font.italic = italic
        return p

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    cpara('AL LAITH GROUP', 32, bold=True, color=NAVY)
    cpara('Weekly Marketing Intelligence Brief', 22, color=GOLD)
    doc.add_paragraph()
    cpara(week_label, 14, color=GREY)
    doc.add_paragraph()
    data_note = '✅ Includes live social media research' if used_live_data else 'Based on AI knowledge base'
    cpara(data_note, 10, color=GREEN if used_live_data else LGREY)
    doc.add_paragraph()
    cpara('Prepared by Al Laith Business OS  ·  Confidential', 10, color=LGREY, italic=True)
    doc.add_page_break()

    # ── Helper: markdown → Word ─────────────────────────────────
    def add_md(text):
        for line in text.split('\n'):
            s = line.strip()
            if not s:
                doc.add_paragraph()
                continue
            if s.startswith('### '):
                h = doc.add_heading(s[4:], level=3)
                for r in h.runs:
                    r.font.color.rgb = NAVY
                    r.font.size = Pt(12)
            elif s.startswith('## '):
                h = doc.add_heading(s[3:], level=2)
                for r in h.runs:
                    r.font.color.rgb = NAVY
                    r.font.size = Pt(15)
            elif s.startswith('# '):
                h = doc.add_heading(s[2:], level=1)
                for r in h.runs:
                    r.font.color.rgb = NAVY
                    r.font.size = Pt(18)
            elif s.startswith(('- ', '• ')):
                p = doc.add_paragraph(s[2:], style='List Bullet')
                if p.runs: p.runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                parts = s.split('**')
                for i, part in enumerate(parts):
                    if part:
                        r = p.add_run(part)
                        r.font.size = Pt(11)
                        r.bold = (i % 2 == 1)

    def add_divider():
        p = doc.add_paragraph()
        r = p.add_run('─' * 70)
        r.font.color.rgb = GOLD
        r.font.size = Pt(7)

    def section_title(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = NAVY
            r.font.size = Pt(18)

    # ── Executive Summary ───────────────────────────────────────
    section_title('Executive Summary')
    p = doc.add_paragraph()
    r = p.add_run(f'Top insights and actions — {week_label}')
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = GREY
    doc.add_paragraph()

    for line in exec_summary.strip().split('\n'):
        s = line.strip()
        if s:
            p = doc.add_paragraph(s, style='List Bullet')
            if p.runs: p.runs[0].font.size = Pt(12)

    add_divider()
    doc.add_page_break()

    # ── Section 1: Competitors ──────────────────────────────────
    section_title('Section 1 — Competitor Social Media & Marketing')
    p = doc.add_paragraph(
        f'Live social media intelligence on {len(competitor_reports)} key UAE competitors. '
        f'Includes LinkedIn, Instagram, and web research gathered {week_label}.'
        if used_live_data else
        f'Marketing intelligence on {len(competitor_reports)} key UAE competitors. {week_label}.'
    )
    if p.runs:
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = GREY
    doc.add_paragraph()

    for comp_name, report_text in competitor_reports:
        add_md(report_text)
        add_divider()
        doc.add_paragraph()

    doc.add_page_break()

    # ── Section 2: UAE Trends ───────────────────────────────────
    section_title('Section 2 — UAE Market Trends')
    add_md(uae_trends)
    doc.add_page_break()

    # ── Section 3: Global B2B ───────────────────────────────────
    section_title('Section 3 — Global B2B Marketing Trends')
    add_md(global_b2b)

    # ── Back cover ──────────────────────────────────────────────
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    cpara('Al Laith Group', 16, bold=True, color=NAVY)
    cpara('Marketing Intelligence Brief', 12, color=GOLD)
    cpara(week_label, 11, color=GREY)
    cpara('Generated by Al Laith Business OS  ·  For internal use only', 9, color=LGREY, italic=True)

    return doc


# ── Main ─────────────────────────────────────────────────────────────────────
def main():
    print(f'\n{"="*60}')
    print('  Al Laith Weekly Marketing Intelligence Report')
    print(f'  {datetime.now().strftime("%A %d %B %Y — %H:%M")}')
    print(f'{"="*60}\n')

    # Check dependencies
    try:
        import anthropic
    except ImportError:
        print('ERROR: anthropic not installed. Run: pip install anthropic\n')
        sys.exit(1)

    # Set up clients
    anthropic_key = get_anthropic_key()
    claude         = anthropic.Anthropic(api_key=anthropic_key)

    tavily_key    = get_tavily_key()
    tavily_client = None
    if tavily_key:
        try:
            from tavily import TavilyClient
            tavily_client = TavilyClient(api_key=tavily_key)
            print('  ✅  Tavily connected — live social media research enabled\n')
        except ImportError:
            print('  ⚠️   tavily-python not installed. Run: pip install tavily-python')
            print('       Continuing without live social data...\n')

    used_live_data = tavily_client is not None
    week_label     = f'Week of {datetime.now().strftime("%d %B %Y")}'
    total          = len(COMPETITORS) + 3

    # ── Step 1: Research competitor social media ────────────────
    competitor_reports = []
    comp_summaries     = ''

    for i, comp in enumerate(COMPETITORS, 1):
        print(f'  [{i}/{total}] {comp["name"]}')

        social_data = ''
        if tavily_client:
            print(f'         → Searching social media & web...')
            social_data = research_competitor_social(tavily_client, comp)
            found = 'Live data found' if social_data else 'No live data — using knowledge base'
            print(f'         → {found}')

        print(f'         → Analysing with Claude...')
        try:
            report = ask(claude, prompt_competitor(comp, social_data), max_tokens=1400)
            competitor_reports.append((comp['name'], report))
            comp_summaries += f'\n{comp["name"]} ({comp["sector"]}): {report[:250]}...\n'
        except Exception as e:
            print(f'         ⚠️  Claude failed: {e}')
            competitor_reports.append((comp['name'], f'## {comp["name"]}\n\nResearch unavailable this week.'))

    # ── Step 2: UAE trends ──────────────────────────────────────
    print(f'\n  [{len(COMPETITORS)+1}/{total}] UAE market trends...')
    uae_live, b2b_live = '', ''
    if tavily_client:
        print('         → Searching UAE market data...')
        uae_live, b2b_live = research_trends_social(tavily_client)

    uae_trends = ask(claude, prompt_uae_trends(uae_live), max_tokens=2000)

    # ── Step 3: Global B2B ──────────────────────────────────────
    print(f'  [{len(COMPETITORS)+2}/{total}] Global B2B marketing trends...')
    global_b2b = ask(claude, prompt_global_b2b(b2b_live), max_tokens=2000)

    # ── Step 4: Executive summary ───────────────────────────────
    print(f'  [{len(COMPETITORS)+3}/{total}] Writing executive summary...')
    exec_summary = ask(
        claude,
        prompt_exec_summary(comp_summaries, uae_trends[:700], global_b2b[:700]),
        max_tokens=400
    )

    # ── Step 5: Build Word doc ──────────────────────────────────
    print('\n  Building Word document...')
    doc = build_document(
        week_label, exec_summary,
        competitor_reports, uae_trends, global_b2b,
        used_live_data
    )

    filename = f'Al Laith Intelligence Brief — {datetime.now().strftime("%Y-%m-%d")}.docx'
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    print(f'\n  ✅  Report saved: {filepath}')
    print(f'  📎  Google Docs: drive.google.com → New → File upload')
    print(f'  💰  Estimated cost this run: ~${"0.12" if used_live_data else "0.08"}')
    print(f'\n{"="*60}\n')


if __name__ == '__main__':
    main()
